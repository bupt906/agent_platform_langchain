"""文档审阅工具函数。

提供文档解析、句子切分、知识库检索等核心能力。
知识库检索通过外部万悟平台 hit 接口完成（见 knowledge_bases/client.py）。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from agent_platform.agents.document_review.knowledge_bases.client import KnowledgeHitClient

# ── 句子切分正则 ──────────────────────────────────────────

_SENTENCE_SPLIT_RE = re.compile(
    r"(?<=[。！？.!?\n])\s*"  # 中英文句末标点 + 换行后切分
)


def parse_document(file_path: str) -> str:
    """根据文件路径或 URL 解析文档，返回全文文本。

    支持格式：txt、md、docx
    支持来源：本地文件路径、HTTP/HTTPS URL
    """
    if file_path.startswith(("http://", "https://")):
        return _parse_from_url(file_path)

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    if suffix in (".txt", ".md", ".markdown"):
        return _parse_text(path)
    elif suffix == ".docx":
        return _parse_docx(path)
    elif suffix == ".doc":
        return _parse_doc(path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}，支持 txt / md / docx / doc")


def _parse_from_url(url: str) -> str:
    """从 HTTP URL 下载文档并解析（带超时，避免慢速源挂死工作线程）。"""
    import tempfile
    import urllib.request
    from urllib.parse import quote, urlparse, urlunparse

    parsed = urlparse(url)
    suffix = Path(parsed.path).suffix.lower()
    # 对路径中的非 ASCII 字符做百分号编码，避免 urllib 抛出 ASCII 编码错误
    encoded_url = urlunparse(parsed._replace(path=quote(parsed.path, safe='/:@')))

    # 下载前先检查 Content-Length（若服务端提供）
    max_size = 50 * 1024 * 1024  # 50MB
    try:
        head_req = urllib.request.Request(encoded_url, method="HEAD")
        with urllib.request.urlopen(head_req, timeout=10) as resp:
            content_length = resp.headers.get("Content-Length")
            if content_length and int(content_length) > max_size:
                raise ValueError(f"文件过大 ({int(content_length)} bytes)，最大允许 {max_size} bytes")
    except urllib.error.URLError:
        pass  # HEAD 请求失败时继续尝试 GET

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        urllib.request.urlretrieve(encoded_url, tmp.name)
        tmp_path = Path(tmp.name)
        if tmp_path.stat().st_size > max_size:
            tmp_path.unlink(missing_ok=True)
            raise ValueError(f"下载文件过大，最大允许 {max_size} bytes")

    try:
        if suffix in (".txt", ".md", ".markdown"):
            return tmp_path.read_text(encoding="utf-8")
        elif suffix == ".docx":
            from docx import Document
            doc = Document(str(tmp_path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif suffix == ".doc":
            return _parse_doc(tmp_path)
        else:
            raise ValueError(f"不支持的 URL 文件格式: {suffix}")
    finally:
        tmp_path.unlink(missing_ok=True)


def _parse_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except ImportError:
        raise ImportError("python-docx 未安装，无法解析 docx 文件。请运行: pip install python-docx")


def _parse_doc(path: Path) -> str:
    """解析 .doc 文件（旧版 Word 格式），优先用 textutil(macOS)/antiword，回退到 python-docx 尝试。"""
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception:
        pass

    import subprocess
    for cmd in (["textutil", "-convert", "txt", "-stdout", str(path)], ["antiword", str(path)]):
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    raise ValueError(f"无法解析 .doc 文件: {path}。请转换为 .docx 或 .txt 格式")


def split_sentences(text: str) -> list[str]:
    """按句子粒度切分文本，过滤空句和纯标点/空白句。

    切分策略：
    - 按 。！？.!?\n 切分
    - 保留分句长度 >= 2 个字符的句子
    - 去除纯标点/数字/空白行
    """
    raw = _SENTENCE_SPLIT_RE.split(text)
    sentences = []
    for s in raw:
        cleaned = s.strip()
        # 过滤空句、纯标点、纯数字、过短的句子
        if not cleaned:
            continue
        if len(cleaned) < 2:
            continue
        if re.match(r"^[\d\s\.\,\;\:\!\?\-—\+\(\)\[\]\{\}]+$", cleaned):
            continue
        sentences.append(cleaned)

    return sentences


async def search_knowledge_bases(
    kb_ids: list[str],
    sentence: str,
    kb_client: KnowledgeHitClient,
) -> list[dict]:
    """调用外部知识库 hit 接口，查找与句子最相关的条文。

    检索参数（matchType / topK / threshold 等）由 settings 统一配置，
    失败重试由 client 内部处理，重试后仍失败则向上抛出异常。

    Args:
        kb_ids: 知识库 ID 列表
        sentence: 待审查句子
        kb_client: 外部知识库客户端

    Returns:
        [{"kb_id": "...", "kb_file": "...", "content": "...", "relevance": 0.8}, ...]
    """
    hits = await kb_client.hit(kb_ids, sentence)
    return [
        {
            "kb_id": h.kb_id,
            "kb_file": h.kb_file,
            "content": h.content,
            "relevance": h.relevance,
        }
        for h in hits
    ]


def format_kb_results_for_prompt(results: list[dict]) -> str:
    """将 KB 检索结果格式化为 LLM prompt 用的上下文文本。"""
    if not results:
        return "（无相关检索结果）"

    parts = []
    for i, r in enumerate(results, 1):
        parts.append(
            f"### 检索结果 {i}（知识库id：{r.get('kb_id', '')}，"
            f"知识库文件：{r.get('kb_file', '')}，相关度：{r.get('relevance', 0)}）\n"
            f"{r.get('content', '')}"
        )

    return "\n\n".join(parts)

